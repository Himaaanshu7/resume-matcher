"""
generate_doc.py
Generates detailed project documentation as a Word (.docx) file.
Run: python generate_doc.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import io

# ── Font paths ────────────────────────────────────────────────────────────────
FONT_DIR = "C:/Windows/Fonts/"
TNR        = FONT_DIR + "times.ttf"
TNR_BOLD   = FONT_DIR + "timesbd.ttf"
TNR_ITALIC = FONT_DIR + "timesi.ttf"
ARIAL      = FONT_DIR + "arial.ttf"
ARIAL_BOLD = FONT_DIR + "arialbd.ttf"
COURIER    = FONT_DIR + "cour.ttf"

def fnt(path, size):
    try:
        return ImageFont.truetype(path, size)
    except Exception:
        return ImageFont.load_default()

# ── Doc helpers ───────────────────────────────────────────────────────────────
doc = Document()

# Set default font to Times New Roman throughout
from docx.oxml.ns import qn
styles_element = doc.styles.element
rPrDefault = styles_element.find('.//' + qn('w:rPrDefault'))

def set_doc_font(doc, font_name="Times New Roman", size_pt=12):
    from docx.oxml import OxmlElement
    styles = doc.styles
    for style in styles:
        try:
            style.font.name = font_name
            style.font.size = Pt(size_pt)
            from docx.oxml.ns import nsmap
        except Exception:
            pass
    # Also set normal style explicitly
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(12)

set_doc_font(doc)

# Set page margins
from docx.oxml import OxmlElement
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_font(cell, font_name="Times New Roman", size=11, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)


def heading(text, level=1, color=(26, 35, 126), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_heading(text, level=level)
    p.alignment = align
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(*color)
    return p


def para(text="", bold=False, italic=False, size=12, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, bold_part="", rest=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part:
        r = p.add_run(bold_part)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(26, 35, 126)
    r2 = p.add_run(rest or text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    return p


def add_table(headers, rows_data, header_bg="1A237E", alt_colors=("F5F5F5", "FFFFFF"), col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_font(hdr_cells[i], bold=True, color=(255, 255, 255), size=11)
    for ri, row_data in enumerate(rows_data):
        row = table.add_row().cells
        for ci, val in enumerate(row_data):
            row[ci].text = val
            set_cell_bg(row[ci], alt_colors[ri % 2])
            set_cell_font(row[ci], size=11)
    if col_widths:
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)
    return table


def img_to_doc(buf, width=6.3):
    doc.add_picture(buf, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def make_img(w, h, bg, fn):
    img = Image.new("RGB", (w, h), bg)
    draw = ImageDraw.Draw(img)
    fn(draw, img)
    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    buf.seek(0)
    return buf


def rounded_box(draw, x1, y1, x2, y2, bg, border, radius=10, width=2):
    draw.rounded_rectangle([x1, y1, x2, y2], radius=radius, fill=bg, outline=border, width=width)


def centered_text(draw, x, y, text, font, color):
    draw.text((x, y), text, fill=color, font=font, anchor="mm")


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
def draw_cover(draw, img):
    W, H = img.size
    # Deep blue gradient background
    for y in range(H):
        ratio = y / H
        r = int(15 + ratio * 25)
        g = int(25 + ratio * 40)
        b = int(100 + ratio * 60)
        draw.line([(0, y), (W, y)], fill=(r, g, b))

    # Decorative horizontal lines
    for yi in [10, 14, H-14, H-10]:
        draw.line([(0, yi), (W, yi)], fill=(255, 255, 255, 60), width=2)

    # White card
    draw.rounded_rectangle([50, 60, W-50, H-60], radius=18,
                            fill=(255, 255, 255), outline="#90CAF9", width=3)

    f_title  = fnt(TNR_BOLD,   44)
    f_sub    = fnt(TNR_ITALIC, 20)
    f_small  = fnt(TNR,        15)
    f_tag    = fnt(ARIAL_BOLD, 13)
    f_url    = fnt(ARIAL,      13)

    draw.text((W//2, 130), "AI Resume Matcher", fill="#0D1B6E",
              font=f_title, anchor="mm")
    draw.line([(180, 160), (W-180, 160)], fill="#90CAF9", width=2)
    draw.text((W//2, 195), "Intelligent Resume & Job Description Analysis Tool",
              fill="#37474F", font=f_sub, anchor="mm")
    draw.text((W//2, 235),
              "Semantic NLP · Tech Keyword Detection · AI Coaching · Free Deployment",
              fill="#546E7A", font=f_small, anchor="mm")

    # Tag pills
    tags = [("Python", "#1565C0"), ("NLP / LLM", "#6A1B9A"),
            ("Streamlit", "#E65100"), ("Groq API", "#2E7D32"), ("Open Source", "#880E4F")]
    total_w = sum(len(t)*9 + 30 for t, _ in tags) + (len(tags)-1)*12
    x = (W - total_w) // 2
    for tag, col in tags:
        tw = len(tag)*9 + 30
        draw.rounded_rectangle([x, 270, x+tw, 300], radius=12, fill=col)
        draw.text((x + tw//2, 285), tag, fill="white", font=f_tag, anchor="mm")
        x += tw + 12

    # Author & URL
    draw.text((W//2, 335), "Himanshu Mishra  ·  himanshu.mishra2810@gmail.com",
              fill="#37474F", font=f_small, anchor="mm")
    draw.text((W//2, 368),
              "Live: https://resume-matcher-pdjd3nkkyhwlxdmfztladr.streamlit.app",
              fill="#1565C0", font=f_url, anchor="mm")
    draw.text((W//2, 393),
              "GitHub: https://github.com/Himaaanshu7/resume-matcher",
              fill="#1565C0", font=f_url, anchor="mm")

cover_buf = make_img(900, 460, "#0D1B6E", draw_cover)
img_to_doc(cover_buf, width=6.3)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1, color=(26, 35, 126), align=WD_ALIGN_PARAGRAPH.CENTER)
toc_items = [
    ("1.", "Project Overview"),
    ("2.", "Key Features"),
    ("3.", "System Architecture"),
    ("4.", "Tech Stack"),
    ("5.", "How the Match Score Works"),
    ("6.", "Tech Keyword Gap Analysis"),
    ("7.", "Per-Bullet Scoring"),
    ("8.", "AI Features — Feedback & Bullet Rewriter"),
    ("9.", "Project File Structure"),
    ("10.", "Deployment"),
    ("11.", "How to Run Locally"),
    ("12.", "How to Use the App"),
    ("13.", "Limitations & Future Improvements"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"  {num}  ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(26, 35, 126)
    r2 = p.add_run(title)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview", 1)
para(
    "AI Resume Matcher is a free, fully open-source web application that helps job seekers "
    "objectively measure how well their resume aligns with a specific job description. "
    "The tool solves a common problem: candidates apply to jobs without knowing which skills, "
    "tools, or keywords are missing from their resume — leading to rejections that could have "
    "been avoided with targeted improvements."
)
para()
para(
    "The application combines two complementary AI techniques: a semantic embedding model "
    "(sentence-transformers) for understanding meaning and context, and a large language model "
    "(Groq's llama-3.1-8b) for generating human-quality coaching advice. Both run efficiently "
    "at zero cost to the user."
)
para()

add_table(
    ["Item", "Details"],
    [
        ("Live Application", "https://resume-matcher-pdjd3nkkyhwlxdmfztladr.streamlit.app/"),
        ("Source Code",      "https://github.com/Himaaanshu7/resume-matcher"),
        ("Author",           "Himanshu Mishra"),
        ("Stack",            "Python, Streamlit, sentence-transformers, Groq API"),
        ("Cost",             "100% free — no subscriptions, no hidden charges"),
        ("Deployment",       "Streamlit Community Cloud"),
    ],
    col_widths=[2.0, 4.3]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Key Features", 1)

features_detailed = [
    (
        "Semantic Match Score (0–100%)",
        "Converts both the resume and job description into 384-dimensional semantic vectors "
        "using the all-MiniLM-L6-v2 sentence-transformer model, then computes cosine similarity "
        "between them. Unlike simple keyword counting, this approach understands meaning — "
        "so 'built REST APIs' and 'developed web services' are recognized as related. "
        "The result is displayed on a colour-coded gauge: green (≥60%), amber (40–60%), red (<40%)."
    ),
    (
        "Tech Keyword Gap Analysis",
        "Scans the job description against a hand-curated vocabulary of 150+ real technologies "
        "across 8 categories: programming languages, web frameworks, cloud/DevOps, databases, "
        "ML/data tools, testing tools, APIs/protocols, and architecture patterns. "
        "A secondary semantic similarity check (threshold 0.55) ensures that aliases like "
        "'sklearn' vs 'scikit-learn' or 'Postgres' vs 'PostgreSQL' are not double-counted. "
        "No soft-skill jargon is ever shown."
    ),
    (
        "Per-Bullet Scoring",
        "Every bullet point extracted from the resume is individually scored against the job "
        "description using the same sentence-transformer model. Results are shown in a sortable "
        "table with three strength tiers: Strong (≥45%), OK (30–45%), Weak (<30%). "
        "This pinpoints exactly which experience statements need improvement."
    ),
    (
        "AI Bullet Rewriter",
        "The three weakest bullets are surfaced with an expandable panel and a one-click "
        "'Rewrite with LLM' button. Groq's llama-3.1-8b-instant model rewrites each bullet "
        "using strong action verbs, JD-aligned terminology, and specific language — "
        "keeping it under 20 words. Results persist on screen via Streamlit session state."
    ),
    (
        "AI Coach Feedback",
        "A single button generates a structured coaching report by sending the resume, "
        "job description, and missing keywords to Groq's LLM. The report contains: "
        "a 2-sentence match summary, 3 specific bullets to add or rewrite with examples, "
        "and 2 critical skills completely absent from the resume. "
        "The result is stored in session state so it doesn't disappear on rerender."
    ),
    (
        "Resume Section Detection",
        "Automatically identifies standard resume sections (Experience, Education, Skills, "
        "Projects, Summary, Certifications, etc.) using regex pattern matching on section headers. "
        "Detected sections are displayed in the results panel."
    ),
    (
        "Multi-Format Resume Upload",
        "Accepts PDF files (via pdfplumber), Microsoft Word DOCX files (via python-docx), "
        "and plain text. Users can also paste resume text directly without uploading a file."
    ),
    (
        "Zero-Friction AI Access",
        "The Groq API key is stored as a Streamlit server secret — users never need to create "
        "an account or enter any credentials. AI features work out of the box for all visitors."
    ),
]

for title, desc in features_detailed:
    bullet(text="", bold_part=f"{title}: ", rest=desc)
    para()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("3. System Architecture", 1)
para(
    "The application is built as four cooperating Python modules, each with a single "
    "responsibility. Streamlit manages the UI and session state, ensuring results persist "
    "across user interactions without page reloads."
)
para()

def draw_arch(draw, img):
    W, H = img.size
    draw.rectangle([0, 0, W, H], fill="#F8FAFF")

    # Draw grid lines lightly
    for x in range(0, W, 80):
        draw.line([(x, 0), (x, H)], fill="#EEF2FF", width=1)
    for y in range(0, H, 60):
        draw.line([(0, y), (W, y)], fill="#EEF2FF", width=1)

    fb = fnt(ARIAL_BOLD, 13)
    fn = fnt(ARIAL, 12)
    fs = fnt(ARIAL, 11)

    # Main boxes
    boxes = [
        # x1, y1, x2, y2, bg, border, title, subtitle
        (30,  120, 155, 200, "#E3F2FD", "#1565C0", "User",           "Browser"),
        (210, 80,  390, 240, "#E8F5E9", "#2E7D32", "app.py",         "Streamlit UI\n+ Session State"),
        (460, 20,  660, 100, "#FFF3E0", "#E65100", "resume_parser.py","PDF/DOCX → Text"),
        (460, 120, 660, 200, "#F3E5F5", "#6A1B9A", "matcher.py",     "Score + Keywords\n+ Bullets"),
        (460, 220, 660, 300, "#FCE4EC", "#880E4F", "llm_feedback.py","Groq API\nFeedback + Rewrite"),
        (210, 280, 390, 360, "#E0F2F1", "#00695C", "Session State",  "Persists results\nacross reruns"),
    ]

    for x1, y1, x2, y2, bg, border, title, sub in boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=10, fill=bg, outline=border, width=2)
        cx = (x1+x2)//2
        cy = (y1+y2)//2
        sub_lines = sub.split("\n")
        offset = -10 * (len(sub_lines))
        draw.text((cx, cy + offset), title, fill=border, font=fb, anchor="mm")
        for i, sl in enumerate(sub_lines):
            draw.text((cx, cy + offset + 20 + i*18), sl, fill="#424242", font=fs, anchor="mm")

    # Arrows
    ac = "#607D8B"
    aw = 2
    # User → app.py
    draw.line([(156, 160), (209, 160)], fill=ac, width=aw)
    draw.polygon([(207,155),(217,160),(207,165)], fill=ac)
    # app.py → resume_parser
    draw.line([(390, 130), (459, 60)], fill=ac, width=aw)
    draw.polygon([(455,55),(465,60),(455,65)], fill=ac)
    # app.py → matcher
    draw.line([(390, 160), (459, 160)], fill=ac, width=aw)
    draw.polygon([(457,155),(467,160),(457,165)], fill=ac)
    # app.py → llm_feedback
    draw.line([(390, 200), (459, 260)], fill=ac, width=aw)
    draw.polygon([(455,255),(465,260),(455,265)], fill=ac)
    # results → session state
    draw.line([(390, 200), (390, 280)], fill=ac, width=aw)
    draw.polygon([(385,278),(390,290),(395,278)], fill=ac)

    # Labels on arrows
    draw.text((183, 148), "input", fill="#78909C", font=fnt(ARIAL,10), anchor="mm")
    draw.text((700, 160), "←  results stored  →", fill="#78909C", font=fn, anchor="mm")

    # Legend
    legend = [("#1565C0","UI Layer"), ("#E65100","Parser"), ("#6A1B9A","ML Engine"),
              ("#880E4F","LLM Layer"), ("#00695C","State")]
    for i, (col, lbl) in enumerate(legend):
        xi = 30 + i*140
        draw.rounded_rectangle([xi, H-45, xi+120, H-15], radius=6, fill=col)
        draw.text((xi+60, H-30), lbl, fill="white", font=fnt(ARIAL_BOLD,11), anchor="mm")

arch_buf = make_img(860, 420, "#F8FAFF", draw_arch)
img_to_doc(arch_buf, width=6.3)
para()
para(
    "Figure 1: System architecture showing the four Python modules and how data flows "
    "from user input through parsing, scoring, and LLM feedback back to the Streamlit UI.",
    italic=True, size=10, color=(96, 125, 139), align=WD_ALIGN_PARAGRAPH.CENTER
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Tech Stack", 1)
para(
    "All libraries used are free and open-source. No paid API is baked into the cost of "
    "running the app for visitors — the Groq free tier provides 14,400 requests per day, "
    "which is more than sufficient for a demo/portfolio application."
)
para()

add_table(
    ["Component", "Library / Service", "Version", "Purpose"],
    [
        ("Semantic Similarity",  "sentence-transformers",     "≥2.7",  "Encode text to vectors, cosine similarity"),
        ("Embedding Model",      "all-MiniLM-L6-v2",          "—",     "80MB model, runs on CPU, 384-dim embeddings"),
        ("Resume Parsing (PDF)", "pdfplumber",                "≥0.11", "Extract text from PDF files page by page"),
        ("Resume Parsing (DOCX)","python-docx",               "≥1.1",  "Extract paragraphs from Word documents"),
        ("LLM Feedback",         "Groq API",                  "≥0.9",  "llama-3.1-8b-instant for coaching & rewrites"),
        ("UI Framework",         "Streamlit",                 "≥1.32", "Web UI, session state, file uploader"),
        ("Visualisation",        "Plotly",                    "≥5.20", "Interactive gauge chart for match score"),
        ("Data Display",         "Pandas",                    "≥2.2",  "Bullet score table rendering"),
        ("ML Utilities",         "PyTorch",                   "≥2.2",  "Tensor operations for embedding similarity"),
        ("Deployment",           "Streamlit Community Cloud", "—",     "Free hosting with secrets management"),
    ],
    col_widths=[1.6, 1.8, 0.7, 2.3]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. SCORING
# ══════════════════════════════════════════════════════════════════════════════
heading("5. How the Match Score Works", 1)
para(
    "The overall match score measures the semantic similarity between the full resume text "
    "and the full job description. It does not rely on keyword counting — instead it uses "
    "deep language understanding so that different words expressing the same concept score highly."
)
para()

def draw_scoring(draw, img):
    W, H = img.size
    draw.rectangle([0, 0, W, H], fill="#FAFAFA")

    fb = fnt(ARIAL_BOLD, 14)
    fn = fnt(ARIAL, 12)
    fs = fnt(ARIAL, 11)

    steps = [
        (40,  "#E3F2FD", "#1565C0", "STEP 1",  "Text Input",
         ["Resume text", "Job Description", "entered by user"]),
        (210, "#E8F5E9", "#2E7D32", "STEP 2",  "Vectorisation",
         ["all-MiniLM-L6-v2", "encodes both texts", "→ 384-dim vectors"]),
        (380, "#FFF3E0", "#E65100", "STEP 3",  "Cosine Similarity",
         ["dot(A,B) / |A||B|", "measures angle", "between vectors"]),
        (550, "#F3E5F5", "#6A1B9A", "STEP 4",  "Score",
         ["Result × 100", "displayed as %", "on gauge chart"]),
    ]

    for x, bg, border, step, title, lines in steps:
        draw.rounded_rectangle([x, 30, x+150, 230], radius=12,
                               fill=bg, outline=border, width=2)
        draw.text((x+75, 60),  step,  fill=border,   font=fb, anchor="mm")
        draw.text((x+75, 90),  title, fill="#212121", font=fn, anchor="mm")
        draw.line([(x+20, 105), (x+130, 105)], fill=border, width=1)
        for i, line in enumerate(lines):
            draw.text((x+75, 130+i*25), line, fill="#424242", font=fs, anchor="mm")

    # Arrows between steps
    for ax in [192, 362, 532]:
        draw.line([(ax, 130), (ax+16, 130)], fill="#90A4AE", width=2)
        draw.polygon([(ax+14,125),(ax+24,130),(ax+14,135)], fill="#90A4AE")

    # Threshold bar
    bar_x, bar_y = 40, 255
    bar_w = 660
    draw.rounded_rectangle([bar_x, bar_y, bar_x+bar_w, bar_y+35], radius=5,
                            fill="#EEEEEE", outline="#BDBDBD", width=1)
    draw.rounded_rectangle([bar_x, bar_y, bar_x+int(bar_w*0.4), bar_y+35], radius=5, fill="#EF9A9A")
    draw.rounded_rectangle([bar_x+int(bar_w*0.4), bar_y, bar_x+int(bar_w*0.6), bar_y+35], radius=0, fill="#FFCC80")
    draw.rounded_rectangle([bar_x+int(bar_w*0.6), bar_y, bar_x+bar_w, bar_y+35], radius=5, fill="#A5D6A7")

    for txt, xp, col in [("Weak  <40%", bar_x+130, "#C62828"),
                          ("Moderate 40–60%", bar_x+330, "#E65100"),
                          ("Strong  ≥60%", bar_x+550, "#2E7D32")]:
        draw.text((xp, bar_y+17), txt, fill=col, font=fnt(ARIAL_BOLD,13), anchor="mm")

scoring_buf = make_img(760, 310, "#FAFAFA", draw_scoring)
img_to_doc(scoring_buf, width=6.3)
para()
para("Figure 2: Four-step pipeline for computing the match score.",
     italic=True, size=10, color=(96,125,139), align=WD_ALIGN_PARAGRAPH.CENTER)
para()

heading("Score Interpretation", 2)
add_table(
    ["Score Range", "Label", "Meaning", "Recommended Action"],
    [
        ("60% – 100%", "Strong Match",    "Resume language closely mirrors the JD",       "Minor tweaks; apply confidently"),
        ("40% – 59%",  "Moderate Match",  "Some overlap but key areas are missing",        "Use AI feedback to identify gaps"),
        ("0%  – 39%",  "Weak Match",      "Resume and JD are semantically misaligned",     "Significant rewrite recommended"),
    ],
    col_widths=[1.2, 1.2, 2.4, 1.6]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. KEYWORD GAP
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Tech Keyword Gap Analysis", 1)
para(
    "The keyword engine was redesigned to show only real technologies and tools — never "
    "soft skills, generic adjectives, or business jargon. It operates in two passes:"
)
para()
bullet("", "Pass 1 — Exact matching: ",
       "The JD is scanned against a vocabulary of 150+ known technologies. "
       "Any term present in the JD but absent from the resume is flagged as a candidate.")
para()
bullet("", "Pass 2 — Semantic deduplication: ",
       "Each candidate is encoded and compared against every sentence in the resume. "
       "If the semantic similarity exceeds 0.55, the term is removed — this handles "
       "synonyms and abbreviations (e.g. Postgres ↔ PostgreSQL, sklearn ↔ scikit-learn).")
para()
para("The vocabulary covers 8 technology categories:")
para()

add_table(
    ["Category", "Examples from Vocabulary"],
    [
        ("Programming Languages",   "Python, JavaScript, TypeScript, Go, Rust, Java, C++, C#, R, SQL, Kotlin, Swift"),
        ("Web Frameworks",          "React, Angular, Vue.js, Next.js, Django, FastAPI, Flask, Spring Boot, Node.js, Rails"),
        ("Cloud & DevOps",          "AWS, Azure, GCP, Docker, Kubernetes, Terraform, Ansible, Jenkins, GitHub Actions, Helm"),
        ("Databases",               "PostgreSQL, MySQL, MongoDB, Redis, Snowflake, BigQuery, Elasticsearch, DynamoDB, Cassandra"),
        ("ML & Data Engineering",   "PyTorch, TensorFlow, scikit-learn, Spark, Airflow, dbt, MLflow, LangChain, Hugging Face"),
        ("Testing & QA",            "Pytest, Selenium, Cypress, Jest, JUnit, TDD, BDD, Unit Testing, Integration Testing"),
        ("APIs & Protocols",        "REST API, GraphQL, gRPC, WebSocket, OAuth, JWT, OpenAPI, Swagger"),
        ("Architecture & Patterns", "Microservices, Serverless, ETL, Data Pipeline, Message Queue, Kafka, RabbitMQ, CI/CD"),
    ],
    col_widths=[2.0, 4.4]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. BULLET SCORING
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Per-Bullet Scoring", 1)
para(
    "While the overall score gives a high-level view, per-bullet scoring shows exactly "
    "which experience statements are strong and which need improvement. This is often more "
    "actionable than the overall score."
)
para()

def draw_bullets(draw, img):
    W, H = img.size
    draw.rectangle([0, 0, W, H], fill="#FFFFFF")

    fb = fnt(ARIAL_BOLD, 13)
    fn = fnt(ARIAL, 12)
    fs = fnt(ARIAL, 11)

    # Header
    draw.rounded_rectangle([0, 0, W, 40], radius=0, fill="#1A237E")
    for x, lbl in [(200,"Resume Bullet"), (530,"Match %"), (660,"Strength")]:
        draw.text((x, 20), lbl, fill="white", font=fb, anchor="mm")

    rows = [
        ("Designed and deployed microservices on AWS ECS using Docker and Terraform",  "71%", "Strong",   "#E8F5E9", "#2E7D32"),
        ("Built ETL pipelines in Python using Apache Airflow and dbt",                 "64%", "Strong",   "#E8F5E9", "#2E7D32"),
        ("Created REST APIs using FastAPI and PostgreSQL",                              "58%", "Strong",   "#E8F5E9", "#2E7D32"),
        ("Analysed data using pandas and matplotlib",                                  "41%", "OK",       "#FFF8E1", "#F57F17"),
        ("Worked on various data projects during internship",                          "29%", "Weak",     "#FFEBEE", "#C62828"),
        ("Helped team with daily tasks and documentation",                             "12%", "Weak",     "#FFEBEE", "#C62828"),
    ]

    for i, (bullet_text, pct, strength, bg, col) in enumerate(rows):
        y = 50 + i*38
        draw.rectangle([0, y, W, y+36], fill=bg)
        draw.line([(0, y+36), (W, y+36)], fill="#E0E0E0", width=1)
        # Bullet text
        display = bullet_text if len(bullet_text) <= 65 else bullet_text[:62]+"..."
        draw.text((30, y+18), display, fill="#212121", font=fs, anchor="lm")
        # Score
        draw.text((530, y+18), pct, fill=col, font=fb, anchor="mm")
        # Badge
        draw.rounded_rectangle([615, y+6, 710, y+30], radius=8, fill=col)
        draw.text((662, y+18), strength, fill="white", font=fb, anchor="mm")

    # Legend
    for x, col, lbl in [(60,"#2E7D32","Strong ≥45%"), (240,"#F57F17","OK 30–44%"),
                        (420,"#C62828","Weak <30%")]:
        draw.rounded_rectangle([x, H-40, x+160, H-10], radius=8, fill=col)
        draw.text((x+80, H-25), lbl, fill="white", font=fb, anchor="mm")

bullets_buf = make_img(760, 340, "#FFFFFF", draw_bullets)
img_to_doc(bullets_buf, width=6.3)
para()
para("Figure 3: Example bullet scoring table showing Strong, OK, and Weak ratings.",
     italic=True, size=10, color=(96,125,139), align=WD_ALIGN_PARAGRAPH.CENTER)
para()
para(
    "The parser extracts bullets using regex patterns that detect •, -, *, and numbered "
    "lists. It also captures long descriptive lines (>40 characters) that are not formally "
    "bulleted. Each extracted bullet is then encoded and compared against the full JD vector "
    "using cosine similarity."
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. AI FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading("8. AI Features", 1)
para(
    "Both AI features are powered by Groq's free API using the llama-3.1-8b-instant model. "
    "Groq runs inference on custom LPU (Language Processing Unit) hardware, making responses "
    "significantly faster than typical cloud LLM APIs. The free tier provides 14,400 requests/day."
)
para()

heading("8.1  AI Coach Feedback", 2)
para(
    "Clicking 'Generate AI Feedback' constructs a structured prompt containing the job "
    "description (first 1,200 characters), the resume (first 1,200 characters), and the "
    "list of missing tech keywords. The model is instructed to respond in a fixed format:"
)
para()
for item in [
    "Match summary — 2 sentences explaining how well the resume fits the role",
    "3 specific bullets to add or rewrite, with concrete examples",
    "2 critical skills from the JD that are completely absent from the resume",
]:
    bullet(text=item)
para()

heading("8.2  AI Bullet Rewriter", 2)
para(
    "The three weakest-scoring bullets are displayed in expandable panels. Each has a "
    "'Rewrite with LLM' button that sends the original bullet and a JD excerpt to the model "
    "with instructions to: start with a strong action verb, be specific and quantified where "
    "possible, and keep the result under 20 words."
)
para()

def draw_rewrite(draw, img):
    W, H = img.size
    draw.rectangle([0, 0, W, H], fill="#F5F5F5")

    fb = fnt(ARIAL_BOLD, 13)
    fn = fnt(ARIAL, 12)
    fs = fnt(ARIAL, 11)

    # Original
    draw.rounded_rectangle([30, 20, W//2-20, H-20], radius=10,
                            fill="#FFEBEE", outline="#C62828", width=2)
    draw.text((W//4, 45), "ORIGINAL (Weak — 18%)", fill="#C62828", font=fb, anchor="mm")
    draw.line([(50, 60), (W//2-40, 60)], fill="#EF9A9A", width=1)
    orig_lines = ["Worked on data analysis", "projects during internship", "using various tools"]
    for i, line in enumerate(orig_lines):
        draw.text((W//4, 90+i*22), line, fill="#424242", font=fn, anchor="mm")

    # Arrow
    draw.line([(W//2, H//2-10), (W//2+30, H//2-10)], fill="#78909C", width=2)
    draw.polygon([(W//2+28, H//2-15),(W//2+40, H//2-10),(W//2+28, H//2-5)], fill="#78909C")
    draw.text((W//2+15, H//2+10), "Groq", fill="#78909C", font=fs, anchor="mm")
    draw.text((W//2+15, H//2+25), "LLM", fill="#78909C", font=fs, anchor="mm")

    # Rewritten
    draw.rounded_rectangle([W//2+50, 20, W-30, H-20], radius=10,
                            fill="#E8F5E9", outline="#2E7D32", width=2)
    draw.text(((W//2+50+W-30)//2, 45), "REWRITTEN (Strong — 67%)",
              fill="#2E7D32", font=fb, anchor="mm")
    draw.line([(W//2+70, 60), (W-50, 60)], fill="#A5D6A7", width=1)
    new_lines = ["Analysed 50K+ customer records", "using Python and pandas to", "identify $200K in revenue gaps"]
    for i, line in enumerate(new_lines):
        draw.text(((W//2+50+W-30)//2, 90+i*22), line, fill="#1B5E20", font=fn, anchor="mm")

rewrite_buf = make_img(700, 200, "#F5F5F5", draw_rewrite)
img_to_doc(rewrite_buf, width=5.5)
para()
para("Figure 4: Example of a weak bullet being rewritten by the LLM.",
     italic=True, size=10, color=(96,125,139), align=WD_ALIGN_PARAGRAPH.CENTER)

heading("8.3  Session State Persistence", 2)
para(
    "A key technical detail: in Streamlit, every button click triggers a full page rerun. "
    "Early versions of the app lost AI results immediately after generation. The fix was to "
    "store all results (match score, bullet scores, AI feedback, rewrites) in "
    "st.session_state — a dictionary that survives reruns. Results now persist until the "
    "user clicks 'Analyze Match' again."
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 9. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("9. Project File Structure", 1)

def draw_files(draw, img):
    W, H = img.size
    draw.rectangle([0, 0, W, H], fill="#1E1E2E")

    # Window chrome
    draw.rounded_rectangle([0, 0, W, H], radius=10, fill="#1E1E2E", outline="#3D3D5C", width=2)
    draw.rounded_rectangle([0, 0, W, 35], radius=10, fill="#2D2D44")
    for cx, col in [(18,"#FF5F56"),(38,"#FFBD2E"),(58,"#27C93F")]:
        draw.ellipse([cx-7, 11, cx+7, 25], fill=col)
    draw.text((W//2, 17), "resume-matcher — file structure", fill="#AAAACC",
              font=fnt(ARIAL_BOLD, 12), anchor="mm")

    lines = [
        ("resume-matcher/",                         "#569CD6", 0),
        ("│",                                       "#555577", 0),
        ("├── app.py",                              "#9CDCFE", 0),
        ("│   # Streamlit UI, routing, session state management", "#6A9955", 20),
        ("│",                                       "#555577", 0),
        ("├── matcher.py",                          "#9CDCFE", 0),
        ("│   # Cosine similarity, per-bullet scoring, tech keyword gap", "#6A9955", 20),
        ("│",                                       "#555577", 0),
        ("├── llm_feedback.py",                     "#9CDCFE", 0),
        ("│   # Groq API: coaching feedback + bullet rewriter", "#6A9955", 20),
        ("│",                                       "#555577", 0),
        ("├── resume_parser.py",                    "#9CDCFE", 0),
        ("│   # PDF / DOCX / TXT text extraction, section detection", "#6A9955", 20),
        ("│",                                       "#555577", 0),
        ("├── requirements.txt",                    "#CE9178", 0),
        ("│   # All Python dependencies",           "#6A9955", 20),
        ("│",                                       "#555577", 0),
        ("├── .streamlit/",                         "#569CD6", 0),
        ("│   └── secrets.toml.example",            "#CE9178", 20),
        ("│       # Template for Groq API key",     "#6A9955", 40),
        ("│",                                       "#555577", 0),
        ("└── .gitignore",                          "#CE9178", 0),
        ("    # Excludes secrets.toml, __pycache__, venv", "#6A9955", 20),
    ]

    fc = fnt(COURIER, 13)
    y = 50
    for text, color, indent in lines:
        draw.text((20 + indent, y), text, fill=color, font=fc)
        y += 20

files_buf = make_img(820, 530, "#1E1E2E", draw_files)
img_to_doc(files_buf, width=6.3)
para()
para("Figure 5: Project file structure with annotations for each file's responsibility.",
     italic=True, size=10, color=(96,125,139), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
heading("10. Deployment", 1)
para(
    "The application is deployed on Streamlit Community Cloud, which provides free hosting "
    "for public Streamlit apps connected to a GitHub repository. Deployments update "
    "automatically whenever code is pushed to the main branch."
)
para()

add_table(
    ["Item", "Details"],
    [
        ("Platform",         "Streamlit Community Cloud (free tier)"),
        ("Live URL",         "https://resume-matcher-pdjd3nkkyhwlxdmfztladr.streamlit.app/"),
        ("GitHub Repository","https://github.com/Himaaanshu7/resume-matcher"),
        ("Branch",           "main"),
        ("Entry Point",      "resume-matcher/app.py"),
        ("Secret Managed",   "GROQ_API_KEY — stored in Streamlit dashboard, never in code"),
        ("Auto-deploy",      "Yes — every push to main triggers a redeploy"),
        ("Cost",             "Free"),
    ],
    col_widths=[2.0, 4.3]
)
para()
heading("Secret Management", 2)
para(
    "The Groq API key is stored as a Streamlit secret in the deployment dashboard — "
    "it is never committed to the GitHub repository. The app reads it at runtime via "
    "st.secrets['GROQ_API_KEY']. If the secret is absent (e.g. local development), "
    "the app falls back to a sidebar input field where users can enter their own key."
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 11. LOCAL SETUP
# ══════════════════════════════════════════════════════════════════════════════
heading("11. How to Run Locally", 1)

def draw_terminal(draw, img):
    W, H = img.size
    draw.rounded_rectangle([0, 0, W, H], radius=10, fill="#0D1117", outline="#30363D", width=2)
    draw.rounded_rectangle([0, 0, W, 32], radius=10, fill="#161B22")
    draw.text((W//2, 16), "Terminal", fill="#8B949E", font=fnt(ARIAL_BOLD,12), anchor="mm")

    fc = fnt(COURIER, 13)
    lines = [
        ("# 1. Clone the repository",                        "#8B949E"),
        ("git clone https://github.com/Himaaanshu7/resume-matcher.git", "#79C0FF"),
        ("cd resume-matcher/resume-matcher",                 "#79C0FF"),
        ("",                                                 "#FFFFFF"),
        ("# 2. Install dependencies",                        "#8B949E"),
        ("pip install -r requirements.txt",                  "#79C0FF"),
        ("",                                                 "#FFFFFF"),
        ("# 3. Add your Groq API key (optional)",            "#8B949E"),
        ("# Create .streamlit/secrets.toml with:",           "#8B949E"),
        ('# GROQ_API_KEY = "gsk_..."',                       "#A5D6A7"),
        ("",                                                 "#FFFFFF"),
        ("# 4. Run the app",                                 "#8B949E"),
        ("streamlit run app.py",                             "#79C0FF"),
        ("",                                                 "#FFFFFF"),
        ("# App available at http://localhost:8501",         "#A5D6A7"),
    ]
    y = 45
    for text, color in lines:
        if text.startswith("#"):
            draw.text((20, y), text, fill=color, font=fc)
        else:
            draw.text((20, y), "$ " if text else "", fill="#A5D6A7", font=fc)
            draw.text((40, y), text, fill=color, font=fc)
        y += 22

terminal_buf = make_img(760, 390, "#0D1117", draw_terminal)
img_to_doc(terminal_buf, width=6.3)
para()
para("Figure 6: Terminal commands to clone, install, and run the app locally.",
     italic=True, size=10, color=(96,125,139), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 12. USAGE GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading("12. How to Use the App", 1)

steps = [
    ("Open the App",
     "Visit the live URL or run locally. The app loads with an empty two-column layout."),
    ("Upload or Paste Resume",
     "Use the left panel to upload a PDF or DOCX file, or paste resume text directly into "
     "the text area. Both options produce identical results."),
    ("Paste the Job Description",
     "Copy the full job description from LinkedIn, Indeed, or any job board and paste it "
     "into the right panel text area."),
    ("Click 'Analyze Match'",
     "The button activates once both inputs are filled. The semantic model runs locally "
     "and results appear within 2–5 seconds (longer on first run while the model loads)."),
    ("Review the Gauge Chart",
     "The match score is shown as a percentage on a colour-coded gauge. Read the summary "
     "line below it for a quick verdict."),
    ("Check Missing Tech Keywords",
     "The left column shows technology terms from the JD not found in the resume. "
     "These are high-priority additions to consider."),
    ("Review Bullet Scores",
     "The bullet score table shows every resume bullet ranked by JD relevance. "
     "Sort mentally by Weak → OK → Strong to find the most impactful items to fix."),
    ("Rewrite Weak Bullets",
     "Expand any weak bullet panel and click 'Rewrite with LLM'. The AI produces an "
     "improved version in under 3 seconds. The result stays on screen persistently."),
    ("Generate AI Feedback",
     "Click 'Generate AI Feedback' at the bottom for a full coaching report. "
     "This gives a holistic view with specific, actionable recommendations."),
]

for i, (step, desc) in enumerate(steps, 1):
    bullet("", f"Step {i} — {step}: ", desc)
    para()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 13. LIMITATIONS & FUTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("13. Limitations & Future Improvements", 1)

heading("Current Limitations", 2)
limitations = [
    ("Score reflects text similarity, not skill matching",
     "A resume full of general tech language can score higher than one with the exact "
     "right skills written differently. The keyword gap and bullet scores provide better signal."),
    ("Groq free tier rate limits",
     "14,400 requests/day is generous for a demo app, but a high-traffic deployment would "
     "need a paid Groq plan or fallback model."),
    ("Tech vocabulary is static",
     "New frameworks and tools require manual additions to the vocabulary list in matcher.py."),
    ("No user accounts or history",
     "Sessions are stateless — analysis results are lost when the browser tab closes."),
]
for title, desc in limitations:
    bullet("", f"{title}: ", desc)
    para()

heading("Potential Future Improvements", 2)
future = [
    "Dynamic tech vocabulary using live job market data APIs",
    "User accounts with analysis history and progress tracking",
    "Side-by-side diff view showing exactly which resume words align with which JD phrases",
    "Export analysis report as a formatted PDF",
    "ATS (Applicant Tracking System) simulation mode",
    "Resume score trend tracking across multiple JD comparisons",
    "Multi-language support for non-English resumes and JDs",
]
for item in future:
    bullet(text=item)

para()
para()
para("— End of Document —", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=(26, 35, 126), size=13)

# ── Save ──────────────────────────────────────────────────────────────────────
output = "Resume_Matcher_Documentation_v2.docx"
doc.save(output)
print(f"Saved: {output}")
